from __future__ import annotations

from pathlib import Path

from docx import Document

from app.document.parser import DocumentParser
from app.document.types import ParsedChapter, ParsedDocument


class DocxParser(DocumentParser):
    def _parse_sync(self, file_path: str | Path) -> ParsedDocument:
        document = Document(str(file_path))
        chapters: list[ParsedChapter] = []
        intro_parts: list[str] = []
        current_title = "Introduction"
        current_content: list[str] = []

        def flush_chapter(title: str, content: list[str]) -> None:
            if not content:
                return
            text = "\n\n".join(part.strip() for part in content if part and part.strip())
            if text:
                chapters.append(ParsedChapter(title=title, content=text))

        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            heading_level = None
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            if "heading" in style_name:
                parts = style_name.split("heading")
                tail = parts[-1].strip().replace(" ", "")
                if tail.isdigit():
                    heading_level = int(tail)
            if heading_level is None:
                for run in paragraph.runs:
                    if run.bold and len(text.split()) <= 12:
                        heading_level = 1
                        break
            if heading_level == 1:
                if current_content:
                    flush_chapter(current_title, current_content)
                current_title = text
                current_content = []
            elif heading_level and heading_level > 1:
                if current_content and current_title != "Introduction":
                    flush_chapter(current_title, current_content)
                    current_title = text
                    current_content = []
                else:
                    intro_parts.append(text)
            else:
                if chapters:
                    current_content.append(text)
                elif current_title == "Introduction":
                    intro_parts.append(text)
                else:
                    current_content.append(text)

        if current_content:
            flush_chapter(current_title, current_content)

        if intro_parts and chapters:
            chapters[0].content = "\n\n".join(intro_parts + [chapters[0].content])
        elif intro_parts and not chapters:
            chapters.append(ParsedChapter(title="Introduction", content="\n\n".join(intro_parts)))

        if not chapters:
            full_text = "\n\n".join((p.text or "").strip() for p in document.paragraphs if (p.text or "").strip())
            chapters.append(ParsedChapter(title="Introduction", content=full_text))

        title = "Document"
        author = None
        metadata = {}
        if document.core_properties.title:
            title = document.core_properties.title.strip() or title
        if document.core_properties.author:
            author = document.core_properties.author.strip() or None
        if document.core_properties.subject:
            metadata["subject"] = document.core_properties.subject

        return ParsedDocument(title=title, author=author, chapters=chapters, metadata=metadata)
