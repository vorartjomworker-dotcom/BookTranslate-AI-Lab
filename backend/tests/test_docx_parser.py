from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from app.document.docx_parser import DocxParser


def _write_docx(path: Path, paragraphs: list[str], heading_indices: set[int] | None = None) -> None:
    document = Document()
    heading_indices = heading_indices or set()
    for index, text in enumerate(paragraphs):
        if index in heading_indices:
            heading = document.add_heading(text, level=1)
            heading.runs[0].bold = True
        else:
            document.add_paragraph(text)
    document.save(path)


def test_docx_parser_handles_two_h1_sections() -> None:
    path = Path("/tmp/test_docx_two_h1.docx")
    _write_docx(
        path,
        [
            "Intro paragraph before the first chapter.",
            "Chapter One",
            "Paragraph in chapter one.",
            "Chapter Two",
            "Paragraph in chapter two.",
        ],
        {1, 3},
    )

    document = asyncio.run(DocxParser().parse(path))

    assert len(document.chapters) == 2
    assert document.chapters[0].title == "Chapter One"
    assert "Paragraph in chapter one." in document.chapters[0].content
    assert document.chapters[1].title == "Chapter Two"


def test_docx_parser_without_headings_uses_introduction() -> None:
    path = Path("/tmp/test_docx_no_headings.docx")
    _write_docx(path, ["Alpha text.", "Beta text.", "Gamma text."])

    document = asyncio.run(DocxParser().parse(path))

    assert len(document.chapters) == 1
    assert document.chapters[0].title == "Introduction"
    assert "Alpha text." in document.chapters[0].content


def test_docx_parser_keeps_intro_before_first_h1() -> None:
    path = Path("/tmp/test_docx_intro_before_h1.docx")
    _write_docx(
        path,
        [
            "Preface sentence.",
            "Second preface sentence.",
            "First Chapter",
            "Later text.",
        ],
        {2},
    )

    document = asyncio.run(DocxParser().parse(path))

    assert len(document.chapters) == 1
    assert document.chapters[0].title == "First Chapter"
    assert "Preface sentence." in document.chapters[0].content or "Second preface sentence." in document.chapters[0].content
