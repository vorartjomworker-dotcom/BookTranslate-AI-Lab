from __future__ import annotations

import io
import json
import stat
import zipfile
from pathlib import Path

import pytest
from unittest.mock import AsyncMock, patch
from fastapi.testclient import TestClient

from app.core.config import settings
from app.core.exceptions import NotFoundError
from app.dependencies.db import get_db
from app.main import app


def _make_docx_bytes() -> bytes:
    from docx import Document

    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Preface sentence before the first chapter.")
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Body of chapter one.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("Body of chapter two.")
    doc.save(buffer)
    return buffer.getvalue()


def _make_epub_bytes() -> bytes:
    from ebooklib import epub

    buffer = io.BytesIO()
    book = epub.EpubBook()
    book.set_identifier("upload-book")
    book.set_title("Upload Book")
    book.set_language("en")
    chapter_one = epub.EpubHtml(title="Chapter One", file_name="chapter1.xhtml", content="<h1>Chapter One</h1><p>Body of chapter one.</p>")
    chapter_two = epub.EpubHtml(title="Chapter Two", file_name="chapter2.xhtml", content="<h1>Chapter Two</h1><p>Body of chapter two.</p>")
    book.add_item(chapter_one)
    book.add_item(chapter_two)
    book.toc = (chapter_one, chapter_two)
    book.spine = [chapter_one, chapter_two]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(buffer, book)
    return buffer.getvalue()


def _make_bad_zip_bytes() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("../evil.txt", "danger")
    return buffer.getvalue()


def _make_zip_with_symlink_or_encrypted_entries(*, symlink: bool = False, encrypted: bool = False) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("safe.txt", "hello")
        if symlink:
            info = zipfile.ZipInfo("link.txt")
            info.create_system = 3
            info.external_attr = 0o120777 << 16
            archive.writestr(info, "target")
        if encrypted:
            info = zipfile.ZipInfo("secret.txt")
            info.flag_bits = 0x1
            archive.writestr(info, "hidden")

    payload = bytearray(buffer.getvalue())
    if encrypted:
        secret_pos = payload.rfind(b"secret.txt")
        local_pos = payload.rfind(b"PK\x03\x04", 0, secret_pos)
        central_pos = payload.rfind(b"PK\x01\x02", 0, secret_pos)
        if local_pos != -1:
            payload[local_pos + 6:local_pos + 8] = (0x0001).to_bytes(2, byteorder="little")
        if central_pos != -1:
            payload[central_pos + 8:central_pos + 10] = (0x0001).to_bytes(2, byteorder="little")

    reopened = zipfile.ZipFile(io.BytesIO(bytes(payload)))
    entries = reopened.infolist()
    if symlink:
        link_entry = next((entry for entry in entries if entry.filename == "link.txt"), None)
        assert link_entry is not None
        assert stat.S_ISLNK((link_entry.external_attr >> 16) & 0xFFFF)
    if encrypted:
        assert any(entry.flag_bits & 0x1 for entry in entries)
    return bytes(payload)


@pytest.fixture
def client(editor_client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> TestClient:
    uploads_dir = tmp_path / "uploads"
    monkeypatch.setattr(settings, "upload_dir", str(uploads_dir))
    app.dependency_overrides[get_db] = lambda: AsyncMock()
    with TestClient(app) as test_client:
        test_client.headers.update(editor_client.headers)
        yield test_client
    app.dependency_overrides.clear()


def test_docx_upload_returns_201_and_counts(client: TestClient) -> None:
    payload = {
        "book": {
            "id": 1,
            "title": "Upload title",
            "author": "Upload author",
            "description": None,
            "file_path": "abc123.docx",
            "file_type": "docx",
            "language": "en",
            "status": "parsed",
        },
        "chapters_count": 2,
        "segments_count": 4,
    }
    with patch("app.api.v1.books.DocumentIngestionService.ingest_upload", new=AsyncMock(return_value=payload)):
        response = client.post(
            "/api/v1/books/upload",
            files={"file": ("sample.docx", _make_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"title": "Upload title", "author": "Upload author", "language": "en"},
        )
    assert response.status_code == 201, response.text
    body = response.json()
    assert body["chapters_count"] == 2
    assert body["segments_count"] == 4
    assert body["book"]["file_path"].endswith(".docx")
    assert body["book"]["file_path"] == body["book"]["file_path"].split("/")[-1]
    assert "C:" not in json.dumps(body)


def test_unsupported_format_returns_415(client: TestClient) -> None:
    response = client.post(
        "/api/v1/books/upload",
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )
    assert response.status_code == 415
    body = response.json()
    assert body["code"] == "unsupported_media_type"
    assert body["request_id"] == response.headers["X-Request-ID"]


def test_oversized_upload_returns_413(client: TestClient) -> None:
    payload = b"A" * (25 * 1024 * 1024 + 1)
    response = client.post(
        "/api/v1/books/upload",
        files={"file": ("large.docx", payload, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 413
    body = response.json()
    assert body["code"] == "payload_too_large"


def test_invalid_zip_content_returns_422(client: TestClient) -> None:
    response = client.post(
        "/api/v1/books/upload",
        files={"file": ("bad.epub", _make_bad_zip_bytes(), "application/epub+zip")},
    )
    assert response.status_code == 422
    body = response.json()
    assert body["code"] == "validation_error"


@pytest.mark.parametrize("zip_kind", ["symlink", "encrypted"])
def test_zip_symlink_and_encrypted_entries_are_rejected(client: TestClient, zip_kind: str) -> None:
    payload = _make_zip_with_symlink_or_encrypted_entries(symlink=zip_kind == "symlink", encrypted=zip_kind == "encrypted")
    before_part_count = len(list(settings.upload_dir_path.glob("*.part")))
    response = client.post(
        "/api/v1/books/upload",
        files={"file": ("bad.epub", payload, "application/epub+zip")},
    )
    assert response.status_code == 422
    body = response.json()
    assert body["code"] == "validation_error"
    assert body["request_id"] == response.headers["X-Request-ID"]
    assert len(list(settings.upload_dir_path.glob("*.part"))) == before_part_count
    assert not any(settings.upload_dir_path.glob("*.epub"))


def test_empty_epub_returns_422(client: TestClient) -> None:
    from ebooklib import epub
    buffer = io.BytesIO()
    book = epub.EpubBook()
    book.set_identifier("empty")
    book.set_title("Empty")
    chapter = epub.EpubHtml(title="Empty chapter", file_name="empty.xhtml", content="<p></p>")
    book.add_item(chapter)
    book.spine = [chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(buffer, book)

    response = client.post(
        "/api/v1/books/upload",
        files={"file": ("empty.epub", buffer.getvalue(), "application/epub+zip")},
    )
    assert response.status_code == 422
    assert response.json()["code"] == "validation_error"


def test_error_envelope_contains_request_id(client: TestClient) -> None:
    with patch("app.api.v1.books.BookService.get_book", new=AsyncMock(side_effect=NotFoundError("book", 999999))):
        response = client.get("/api/v1/books/999999")
    assert response.status_code == 404
    body = response.json()
    assert body["request_id"] == response.headers["X-Request-ID"]
    assert body["code"] == "not_found"
